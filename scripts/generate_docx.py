import os
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def generate_word_document():
    print("[DOCX] Generating TECHNICAL_REPORT.docx from TECHNICAL_REPORT.md...")
    
    md_path = 'TECHNICAL_REPORT.md'
    docx_path = 'TECHNICAL_REPORT.docx'
    
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        sys.exit(1)
        
    doc = Document()
    
    # Configure page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)
        
    # Read markdown contents
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    # Set default styles
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    in_bullet = False
    
    for line in lines:
        line_str = line.strip()
        
        # Skip empty lines
        if not line_str:
            continue
            
        # Divider lines
        if line_str == '---':
            p = doc.add_paragraph()
            p_border = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/></w:pBdr>')
            p._p.get_or_add_pPr().append(p_border)
            continue
            
        # Headers
        if line_str.startswith('# '):
            h = doc.add_heading(level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = h.add_run(line_str[2:])
            run.font.name = 'Arial'
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
            # Add some spacing below header
            h.paragraph_format.space_after = Pt(18)
            h.paragraph_format.space_before = Pt(12)
            continue
            
        if line_str.startswith('## '):
            h = doc.add_heading(level=2)
            run = h.add_run(line_str[3:])
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5) # Indigo Accent
            h.paragraph_format.space_before = Pt(16)
            h.paragraph_format.space_after = Pt(8)
            continue
            
        if line_str.startswith('### '):
            h = doc.add_heading(level=3)
            run = h.add_run(line_str[4:])
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)
            continue
            
        # Bullet list items
        if line_str.startswith('- ') or line_str.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            clean_text = line_str[2:]
            
            # Format inline bold text (e.g. **text**)
            parts = clean_text.split('**')
            for idx, part in enumerate(parts):
                run = p.add_run(part)
                if idx % 2 == 1:
                    run.bold = True
            p.paragraph_format.space_after = Pt(4)
            continue
            
        # Numbered list items
        if line_str.startswith('1. ') or line_str.startswith('2. ') or line_str.startswith('3. '):
            p = doc.add_paragraph(style='List Number')
            clean_text = line_str[3:]
            
            parts = clean_text.split('**')
            for idx, part in enumerate(parts):
                run = p.add_run(part)
                if idx % 2 == 1:
                    run.bold = True
            p.paragraph_format.space_after = Pt(4)
            continue
            
        # Normal paragraphs (handling math blocks or code block representations slightly)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15
        
        # Simple inline formatting for **bold** and `code`
        parts = line_str.split('**')
        for idx, part in enumerate(parts):
            # Inner check for code ticks
            subparts = part.split('`')
            for s_idx, subpart in enumerate(subparts):
                run = p.add_run(subpart)
                if idx % 2 == 1:
                    run.bold = True
                if s_idx % 2 == 1:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x8B, 0x5C, 0xF6)
                    
    # Save the output file
    doc.save(docx_path)
    print(f"[DOCX] Success! Formatted Word document saved to {docx_path}")

if __name__ == "__main__":
    generate_word_document()
